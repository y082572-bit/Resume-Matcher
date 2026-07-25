"""Service-level tests for Explicit Provenance Stage P6-B2A-R1: the
production composition root, ``generate_current_deterministic_docx_proposal``
(``cv_docx_proposal_orchestration.py``).

Every test here calls only the production callable -- never
``DeterministicDocxTemplateProvider``, ``DeterministicCvDocxRenderer``,
``CvDocxHeaderLines``, or ``CvDocxRenderPlan`` directly. Header resolution
runs against a real, file-backed SQLite ``CandidateIdentityHeaderSqlService``
(same pattern as ``tests/unit/test_cv_docx_header_binding_builder.py``), and
the P6-A pipeline fixture (``build_ready_integrated_context``/
``build_document_input`` from ``tests.unit.test_cv_document_input_validation``)
supplies a real, fresh, READY ``ApprovedContentDocumentInput`` -- P6-A's own
``validate_approved_content_document_input`` is exercised for real, not
bypassed. The ``TruthEntity``/``Resume``/``CandidateIdentityBinding`` rows
are inserted using the exact ``person_entity_id`` the P6-A pipeline fixture
generates, so the two identities line up the way the real production
``JobArtifactOwnerKey`` -> ``candidate_identity_bindings`` -> Master Resume
-> ``PersonalInfo`` chain requires.
"""

from __future__ import annotations

import hashlib
import zipfile
from io import BytesIO
from pathlib import Path
from uuid import UUID, uuid4

import docx
import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.db_engine import init_models_sync
from app.models import Resume, TruthEntity
from app.schemas.cv_document_artifact import ConfirmedCvPdfArtifact, DocumentProvenanceMode
from app.schemas.cv_docx_visual_contract import DEFAULT_VISUAL_CONTRACT
from app.services.candidate_identity_binding_sql_repository import CandidateIdentityBindingSqlService
from app.services.cv_docx_header_binding_builder import CandidateIdentityHeaderSqlService
from app.services.cv_docx_proposal_orchestration import (
    DeterministicDocxProposalGenerationStatus,
    generate_current_deterministic_docx_proposal,
)
from app.services.cv_document_repository_protocol import CasReplaceStatus, InMemoryCvDocumentArtifactRepository
from tests.unit.test_cv_document_input_validation import build_document_input, build_ready_integrated_context


RENDERING_POLICY_VERSION = "cv-docx-rendering-policy-v1"

_FULL_PERSONAL_INFO = {
    "name": "Marek Grabowski",
    "email": "marek@example.com",
    "phone": "+48 000 000 000",
    "location": "Warsaw, Poland",
    "linkedin": "https://www.linkedin.com/in/marekgrabowski",
    "website": "https://marekgrabowski.com",
}


def _fake_sha256(seed: str) -> str:
    return hashlib.sha256(seed.encode()).hexdigest()


def _document_xml(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(BytesIO(docx_bytes)) as zf:
        return zf.read("word/document.xml").decode("utf-8")


def _full_text(docx_bytes: bytes) -> str:
    document = docx.Document(BytesIO(docx_bytes))
    return "\n".join(p.text for p in document.paragraphs)


@pytest.fixture
def sql_env(tmp_path: Path) -> sessionmaker:
    engine = create_engine(f"sqlite:///{tmp_path / 'p6b2a_r1.db'}", future=True)
    init_models_sync(engine)
    return sessionmaker(engine, expire_on_commit=False)


def _insert_person_entity(session_factory: sessionmaker, person_entity_id: UUID) -> None:
    with session_factory() as session:
        with session.begin():
            session.add(
                TruthEntity(
                    entity_id=str(person_entity_id),
                    entity_type="PERSON",
                    display_name="Test Person",
                    status="ACTIVE",
                    content_fingerprint=_fake_sha256(str(person_entity_id)),
                )
            )


def _insert_master_resume(session_factory: sessionmaker, *, personal_info: dict) -> str:
    resume_id = str(uuid4())
    with session_factory() as session:
        with session.begin():
            session.add(
                Resume(
                    resume_id=resume_id,
                    content="# Resume",
                    is_master=True,
                    processed_data={"personalInfo": personal_info},
                )
            )
    return resume_id


def _update_resume_personal_info(session_factory: sessionmaker, resume_id: str, personal_info: dict) -> None:
    with session_factory() as session:
        with session.begin():
            resume = session.get(Resume, resume_id)
            resume.processed_data = {"personalInfo": personal_info}


def _bind_person_to_resume(
    session_factory: sessionmaker, *, person_entity_id: UUID, personal_info: dict
) -> str:
    _insert_person_entity(session_factory, person_entity_id)
    resume_id = _insert_master_resume(session_factory, personal_info=personal_info)
    binding_service = CandidateIdentityBindingSqlService(session_factory)
    binding_result = binding_service.create_binding(person_entity_id=person_entity_id, master_resume_id=resume_id)
    assert binding_result.status.value == "CREATED"
    return resume_id


async def _bound_document_input(session_factory: sessionmaker, *, personal_info: dict = _FULL_PERSONAL_INFO):
    """Build a real, fresh P6-A ``ApprovedContentDocumentInput`` and bind
    its exact ``person_entity_id`` to a fresh master ``Resume`` carrying
    ``personal_info``, so owner-bound header resolution can find it."""

    ctx = await build_ready_integrated_context()
    document_input = build_document_input(ctx)
    resume_id = _bind_person_to_resume(
        session_factory, person_entity_id=ctx["plan"].person_entity_id, personal_info=personal_info
    )
    return document_input, resume_id


@pytest.mark.asyncio
async def test_owner_bound_header_is_resolved_and_full_header_rendered(sql_env: sessionmaker) -> None:
    document_input, _ = await _bound_document_input(sql_env)
    header_service = CandidateIdentityHeaderSqlService(sql_env)
    repository = InMemoryCvDocumentArtifactRepository()

    result = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )

    assert result.status == DeterministicDocxProposalGenerationStatus.GENERATED
    assert result.artifact is not None
    assert result.docx_bytes is not None

    full_text = _full_text(result.docx_bytes)
    assert "Marek Grabowski" in full_text  # full_name
    assert "marek@example.com" in full_text  # email
    assert "+48 000 000 000" in full_text  # phone
    assert "Warsaw, Poland" in full_text  # location
    assert "linkedin.com/in/marekgrabowski" in full_text  # linkedin, slug preserved

    document_xml = _document_xml(result.docx_bytes)
    assert "marekgrabowski" in document_xml.lower()  # slug not blanket-filtered
    assert "marekgrabowski.com" not in document_xml.lower()  # website domain never rendered


@pytest.mark.asyncio
async def test_approved_content_is_present_in_generated_docx(sql_env: sessionmaker) -> None:
    document_input, _ = await _bound_document_input(sql_env)
    header_service = CandidateIdentityHeaderSqlService(sql_env)
    repository = InMemoryCvDocumentArtifactRepository()

    result = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert result.status == DeterministicDocxProposalGenerationStatus.GENERATED

    full_text = _full_text(result.docx_bytes)
    approved = document_input.approved_content_result
    for section in approved.draft.sections:
        elements = (
            section.elements
            if hasattr(section, "elements")
            else [
                e
                for entry in section.experience_entries
                for e in (entry.header_elements + entry.responsibility_elements + entry.achievement_elements)
            ]
        )
        for element in elements:
            assert element.content_text in full_text


@pytest.mark.asyncio
async def test_exact_renderer_bytes_persist_via_existing_repository(sql_env: sessionmaker) -> None:
    document_input, _ = await _bound_document_input(sql_env)
    header_service = CandidateIdentityHeaderSqlService(sql_env)
    repository = InMemoryCvDocumentArtifactRepository()

    result = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert result.status == DeterministicDocxProposalGenerationStatus.GENERATED

    stored = repository.get_current_proposal(result.artifact.owner_key)
    assert stored == result.artifact
    stored_bytes = repository.read_current_proposal_bytes(result.artifact.owner_key)
    assert stored_bytes == result.docx_bytes
    assert result.artifact.generated_docx_content_hash == hashlib.sha256(result.docx_bytes).hexdigest()
    assert hashlib.sha256(stored_bytes).hexdigest() == result.artifact.generated_docx_content_hash


@pytest.mark.asyncio
async def test_current_proposal_revision_increments_and_stale_revision_is_fail_closed(
    sql_env: sessionmaker,
) -> None:
    document_input, _ = await _bound_document_input(sql_env)
    header_service = CandidateIdentityHeaderSqlService(sql_env)
    repository = InMemoryCvDocumentArtifactRepository()

    first = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert first.status == DeterministicDocxProposalGenerationStatus.GENERATED
    assert first.artifact.proposal_revision == 1

    second = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=1,
    )
    assert second.status == DeterministicDocxProposalGenerationStatus.GENERATED
    assert second.artifact.proposal_revision == 2

    # a stale caller still believing revision is None must be rejected,
    # fail-closed, without touching the current slot -- and via a closed
    # orchestration status, never a raw ProposalBuildStatus.
    stale = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert stale.status == DeterministicDocxProposalGenerationStatus.PROPOSAL_BUILD_FAILED
    assert stale.artifact is None
    assert stale.docx_bytes is None

    current = repository.get_current_proposal(first.artifact.owner_key)
    assert current.proposal_revision == 2
    assert current.artifact_fingerprint == second.artifact.artifact_fingerprint


@pytest.mark.asyncio
async def test_confirmed_pdf_slot_untouched_and_master_resume_unmodified(sql_env: sessionmaker) -> None:
    document_input, resume_id = await _bound_document_input(sql_env)
    header_service = CandidateIdentityHeaderSqlService(sql_env)
    repository = InMemoryCvDocumentArtifactRepository()

    first = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert first.status == DeterministicDocxProposalGenerationStatus.GENERATED
    owner_key = first.artifact.owner_key

    fake_pdf = ConfirmedCvPdfArtifact(
        artifact_id=uuid4(),
        owner_key=owner_key,
        source_validated_docx_snapshot_fingerprint="a" * 64,
        source_docx_sha256="b" * 64,
        approved_cv_content_fingerprint=document_input.approved_content_result.result_fingerprint,
        content_plan_fingerprint=document_input.source_content_plan.content_plan_fingerprint,
        pdf_sha256="c" * 64,
        conversion_adapter_id="fake-pdf-adapter",
        conversion_policy_version="pdf-policy-v1",
        provenance_mode=DocumentProvenanceMode.PIPELINE_UNMODIFIED_DOCUMENT,
        artifact_fingerprint="d" * 64,
        pdf_revision=1,
        superseded=False,
    )
    replace_result = repository.replace_current_pdf(owner_key, fake_pdf, b"%PDF-fake-bytes", None)
    assert replace_result.status == CasReplaceStatus.REPLACED

    second = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=1,
    )
    assert second.status == DeterministicDocxProposalGenerationStatus.GENERATED
    assert repository.get_current_pdf(owner_key).artifact_fingerprint == "d" * 64

    with sql_env() as session:
        resume = session.get(Resume, resume_id)
        assert resume.processed_data["personalInfo"]["name"] == "Marek Grabowski"


@pytest.mark.asyncio
async def test_header_identity_flows_into_proposal_fingerprint_via_linkedin_change(
    sql_env: sessionmaker,
) -> None:
    """A LinkedIn-only change flows through source_personal_info_fingerprint
    -> render_plan_fingerprint -> DOCX bytes -> generated_docx_content_hash
    -> generation_input_fingerprint -> artifact_fingerprint. Two *fresh*
    repositories are used so both builds land on proposal_revision=1 --
    proposal_revision is itself part of artifact_fingerprint, so comparing
    across different revisions would always differ regardless of header
    content."""

    ctx = await build_ready_integrated_context()
    document_input = build_document_input(ctx)
    resume_id = _bind_person_to_resume(
        sql_env, person_entity_id=ctx["plan"].person_entity_id, personal_info=dict(_FULL_PERSONAL_INFO)
    )
    header_service = CandidateIdentityHeaderSqlService(sql_env)

    result_a = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=InMemoryCvDocumentArtifactRepository(),
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert result_a.status == DeterministicDocxProposalGenerationStatus.GENERATED
    assert result_a.artifact.proposal_revision == 1

    personal_info_b = dict(_FULL_PERSONAL_INFO, linkedin="https://www.linkedin.com/in/someone-else")
    _update_resume_personal_info(sql_env, resume_id, personal_info_b)

    result_b = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=InMemoryCvDocumentArtifactRepository(),
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert result_b.status == DeterministicDocxProposalGenerationStatus.GENERATED
    assert result_b.artifact.proposal_revision == 1

    assert result_a.docx_bytes != result_b.docx_bytes
    assert result_a.artifact.generated_docx_content_hash != result_b.artifact.generated_docx_content_hash
    assert result_a.artifact.generation_input_fingerprint != result_b.artifact.generation_input_fingerprint
    assert result_a.artifact.artifact_fingerprint != result_b.artifact.artifact_fingerprint


@pytest.mark.asyncio
async def test_website_only_change_does_not_change_artifact_fingerprint(sql_env: sessionmaker) -> None:
    ctx = await build_ready_integrated_context()
    document_input = build_document_input(ctx)
    resume_id = _bind_person_to_resume(
        sql_env, person_entity_id=ctx["plan"].person_entity_id, personal_info=dict(_FULL_PERSONAL_INFO)
    )
    header_service = CandidateIdentityHeaderSqlService(sql_env)

    result_a = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=InMemoryCvDocumentArtifactRepository(),
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert result_a.status == DeterministicDocxProposalGenerationStatus.GENERATED
    assert result_a.artifact.proposal_revision == 1

    personal_info_b = dict(_FULL_PERSONAL_INFO, website="https://a-totally-different-domain.example")
    _update_resume_personal_info(sql_env, resume_id, personal_info_b)

    result_b = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=InMemoryCvDocumentArtifactRepository(),
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert result_b.status == DeterministicDocxProposalGenerationStatus.GENERATED
    assert result_b.artifact.proposal_revision == 1

    assert result_a.docx_bytes == result_b.docx_bytes
    assert result_a.artifact.generated_docx_content_hash == result_b.artifact.generated_docx_content_hash
    assert result_a.artifact.generation_input_fingerprint == result_b.artifact.generation_input_fingerprint
    assert result_a.artifact.artifact_fingerprint == result_b.artifact.artifact_fingerprint


@pytest.mark.asyncio
async def test_header_resolution_failure_is_closed(sql_env: sessionmaker) -> None:
    ctx = await build_ready_integrated_context()
    document_input = build_document_input(ctx)
    # no TruthEntity/Resume/CandidateIdentityBinding row is ever inserted
    # for this person -- resolve_header_binding must fail closed.
    header_service = CandidateIdentityHeaderSqlService(sql_env)
    repository = InMemoryCvDocumentArtifactRepository()

    result = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert result.status == DeterministicDocxProposalGenerationStatus.HEADER_RESOLUTION_FAILED
    assert result.artifact is None
    assert result.docx_bytes is None


@pytest.mark.asyncio
async def test_owner_mismatch_is_closed(sql_env: sessionmaker) -> None:
    """Binds person A to a resume, then generates a proposal for an
    unrelated person B whose document input carries a different
    ``person_entity_id`` that never appears in ``candidate_identity_bindings``
    -- exercising the same fail-closed owner-identity boundary an actual
    ``BINDING_OWNER_MISMATCH``/tampered-owner-key would hit."""

    ctx_bound = await build_ready_integrated_context()
    _bind_person_to_resume(
        sql_env, person_entity_id=ctx_bound["plan"].person_entity_id, personal_info=_FULL_PERSONAL_INFO
    )

    ctx_unbound = await build_ready_integrated_context()
    document_input_unbound = build_document_input(ctx_unbound)

    header_service = CandidateIdentityHeaderSqlService(sql_env)
    repository = InMemoryCvDocumentArtifactRepository()

    result = generate_current_deterministic_docx_proposal(
        document_input=document_input_unbound,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert result.status in (
        DeterministicDocxProposalGenerationStatus.HEADER_RESOLUTION_FAILED,
        DeterministicDocxProposalGenerationStatus.HEADER_OWNER_MISMATCH,
    )
    assert result.artifact is None
    assert result.docx_bytes is None

    # the shared repository's proposal slot for the *bound* owner is
    # untouched by the mismatched attempt above -- the bound owner's own
    # generation still succeeds independently against the same repository.
    bound_result = generate_current_deterministic_docx_proposal(
        document_input=build_document_input(ctx_bound),
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert bound_result.status == DeterministicDocxProposalGenerationStatus.GENERATED


@pytest.mark.asyncio
async def test_generated_proposal_is_valid_opc_package(sql_env: sessionmaker) -> None:
    document_input, _ = await _bound_document_input(sql_env)
    header_service = CandidateIdentityHeaderSqlService(sql_env)
    repository = InMemoryCvDocumentArtifactRepository()

    result = generate_current_deterministic_docx_proposal(
        document_input=document_input,
        visual_contract=DEFAULT_VISUAL_CONTRACT,
        header_service=header_service,
        repository=repository,
        rendering_policy_version=RENDERING_POLICY_VERSION,
        expected_previous_revision=None,
    )
    assert result.status == DeterministicDocxProposalGenerationStatus.GENERATED
    document = docx.Document(BytesIO(result.docx_bytes))
    assert len(document.paragraphs) > 0
