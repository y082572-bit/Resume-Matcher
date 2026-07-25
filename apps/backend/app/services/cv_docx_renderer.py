"""Explicit Provenance Stage P6-B2a (R3, P6-B2A-R1): the deterministic
Proposal DOCX renderer.

``DeterministicCvDocxRenderer`` implements ``CvDocxRenderingAdapter``
(``cv_document_adapters.py``). It never calls an LLM, never paraphrases,
never adds or removes a claim, and never changes a number -- it places the
blocks of an already-built, complete ``CvDocxRenderPlan`` into the blank
template it receives, in the exact order the plan provides them, and
nothing else.

R3 (P6-B2A-R1): the renderer never builds a render plan itself and never
defaults to ``header_lines=None``. Composition (owner-bound header
resolution, render-plan construction) is the exclusive responsibility of
the caller (``cv_docx_proposal_orchestration.py``) -- this renderer is
handed an already-complete ``CvDocxRenderPlan`` at construction time and
only ever executes it: template application, canonical DOCX packaging,
exact bytes. It never fetches header data, never calls a database or
``CandidateIdentityHeaderSqlService``, never calls an LLM, and never builds
provenance.
"""

from __future__ import annotations

import hashlib
from io import BytesIO

import docx
from docx.document import Document as DocxDocument
from docx.enum.text import WD_BREAK

from app.schemas.cv_content_approval import ApprovalAssemblyStatus, ApprovedCvContentResult
from app.schemas.cv_content_plan import CvContentPlan
from app.schemas.cv_document_artifact import (
    CvDocxRenderingResult,
    DocxRenderingFailureCode,
    DocxRenderingStatus,
)
from app.schemas.cv_docx_render_plan import CvDocxRenderBlock, CvDocxRenderBlockKind, CvDocxRenderPlan
from app.schemas.cv_docx_visual_contract import CvDocxVisualContract, DEFAULT_VISUAL_CONTRACT
from app.services.cv_docx_template import serialize_canonical_document


RENDERER_ADAPTER_ID = "cv-docx-deterministic-renderer-v1"
RENDERING_POLICY_VERSION = "cv-docx-rendering-policy-v1"


def _write_blocks(
    document: DocxDocument, blocks: tuple[CvDocxRenderBlock, ...], visual_contract: CvDocxVisualContract
) -> None:
    for block in blocks:
        if block.kind == CvDocxRenderBlockKind.PAGE_BREAK:
            paragraph = document.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            continue
        if block.kind == CvDocxRenderBlockKind.NAME_HEADER:
            document.add_paragraph(block.text, style="CvName")
            continue
        if block.kind == CvDocxRenderBlockKind.CONTACT_LINE:
            document.add_paragraph(block.text, style="CvContactLine")
            continue
        if block.kind == CvDocxRenderBlockKind.SECTION_HEADER:
            document.add_paragraph(block.text, style="CvSectionHeader")
            continue
        if block.kind == CvDocxRenderBlockKind.BULLET_ITEM:
            # A native Word bullet glyph + tab is used as the marker --
            # never a literal "-"/"- " prefix, so a hyphen-before-bullet
            # can never originate in this renderer.
            document.add_paragraph(f"{visual_contract.bullet_glyph}\t{block.text}", style="CvBullet")
            continue
        # BODY_PARAGRAPH
        document.add_paragraph(block.text, style="Normal")


class DeterministicCvDocxRenderer:
    """Implements ``CvDocxRenderingAdapter``. Stateless and pure aside
    from the fixed ``visual_contract``/``render_plan`` it was constructed
    with -- it must always be paired with a ``DocxTemplateProvider`` built
    from the exact same ``CvDocxVisualContract`` instance/value, since it
    only ever *uses* the named paragraph styles already defined on the
    template it is handed; it never redefines them.

    R3: ``render_plan`` is a required, already-complete
    ``CvDocxRenderPlan`` -- this renderer never builds one itself and
    never falls back to an implicit ``header_lines=None``."""

    def __init__(
        self, visual_contract: CvDocxVisualContract | None = None, *, render_plan: CvDocxRenderPlan
    ) -> None:
        self._visual_contract = visual_contract or DEFAULT_VISUAL_CONTRACT
        self._render_plan = render_plan

    @property
    def renderer_adapter_id(self) -> str:
        return RENDERER_ADAPTER_ID

    def render(
        self,
        *,
        approved_content_result: ApprovedCvContentResult,
        content_plan: CvContentPlan,
        template_bytes: bytes,
        template_fingerprint: str,
        rendering_policy_version: str,
    ) -> CvDocxRenderingResult:
        if (
            approved_content_result.status != ApprovalAssemblyStatus.READY
            or approved_content_result.draft is None
        ):
            return CvDocxRenderingResult(
                status=DocxRenderingStatus.FAILED,
                renderer_adapter_id=self.renderer_adapter_id,
                rendering_policy_version=rendering_policy_version,
                failure_code=DocxRenderingFailureCode.CONTENT_INCOMPATIBLE,
                diagnostics=("approved_content_result is not READY / carries no draft",),
            )
        if approved_content_result.content_plan_fingerprint != content_plan.content_plan_fingerprint:
            return CvDocxRenderingResult(
                status=DocxRenderingStatus.FAILED,
                renderer_adapter_id=self.renderer_adapter_id,
                rendering_policy_version=rendering_policy_version,
                failure_code=DocxRenderingFailureCode.CONTENT_INCOMPATIBLE,
                diagnostics=(
                    "approved_content_result.content_plan_fingerprint does not match content_plan",
                ),
            )
        if self._render_plan.approved_cv_content_fingerprint != approved_content_result.result_fingerprint:
            return CvDocxRenderingResult(
                status=DocxRenderingStatus.FAILED,
                renderer_adapter_id=self.renderer_adapter_id,
                rendering_policy_version=rendering_policy_version,
                failure_code=DocxRenderingFailureCode.CONTENT_INCOMPATIBLE,
                diagnostics=("render_plan.approved_cv_content_fingerprint does not match approved_content_result",),
            )
        if self._render_plan.content_plan_fingerprint != content_plan.content_plan_fingerprint:
            return CvDocxRenderingResult(
                status=DocxRenderingStatus.FAILED,
                renderer_adapter_id=self.renderer_adapter_id,
                rendering_policy_version=rendering_policy_version,
                failure_code=DocxRenderingFailureCode.CONTENT_INCOMPATIBLE,
                diagnostics=("render_plan.content_plan_fingerprint does not match content_plan",),
            )

        try:
            document = docx.Document(BytesIO(template_bytes))
            _write_blocks(document, self._render_plan.blocks, self._visual_contract)
            final_bytes = serialize_canonical_document(document)
        except Exception:
            return CvDocxRenderingResult(
                status=DocxRenderingStatus.FAILED,
                renderer_adapter_id=self.renderer_adapter_id,
                rendering_policy_version=rendering_policy_version,
                failure_code=DocxRenderingFailureCode.RENDERER_ERROR,
                diagnostics=("python-docx failed to open the template or write content",),
            )

        output_sha256 = hashlib.sha256(final_bytes).hexdigest()
        return CvDocxRenderingResult(
            status=DocxRenderingStatus.SUCCEEDED,
            renderer_adapter_id=self.renderer_adapter_id,
            rendering_policy_version=rendering_policy_version,
            docx_bytes=final_bytes,
            output_sha256=output_sha256,
        )
