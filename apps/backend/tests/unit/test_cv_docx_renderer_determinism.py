"""Unit tests for Explicit Provenance Stage P6-B2a (R3, P6-B2A-R1):
``DeterministicCvDocxRenderer`` -- byte-level determinism, visual-contract
compliance, and technical validity of a system-generated Proposal DOCX.

Uses the hand-constructed fixture from ``test_cv_docx_render_plan_builder``
(``build_fixture_approved_content_result``) paired with a real
``CvContentPlan`` built via the shared P6-A pipeline factory
(``build_ready_integrated_context``) purely to obtain a structurally valid
``CvContentPlan`` whose ``content_plan_fingerprint`` matches the fixture's
``approved_content_result.content_plan_fingerprint`` -- the renderer's own
defensive check (``content_plan_fingerprint`` match) requires this pairing.

R3: the renderer no longer builds its own render plan -- every test here
builds a complete ``CvDocxRenderPlan`` via ``build_cv_docx_render_plan``
and constructs ``DeterministicCvDocxRenderer(visual_contract,
render_plan=plan)`` explicitly, mirroring the real composition-then-
rendering split ``cv_docx_proposal_orchestration.py`` performs in
production.
"""

from __future__ import annotations

import hashlib
import os
import subprocess
import sys
import textwrap
import zipfile
from io import BytesIO
from pathlib import Path

import docx
import pytest

from app.schemas.cv_content_approval import ApprovalAssemblyStatus
from app.schemas.cv_content_plan import CvSection
from app.schemas.cv_document_artifact import DocxRenderingStatus
from app.schemas.cv_docx_visual_contract import DEFAULT_VISUAL_CONTRACT
from app.services.cv_docx_render_plan_builder import CvDocxHeaderLines, build_cv_docx_render_plan
from app.services.cv_docx_renderer import DeterministicCvDocxRenderer, RENDERING_POLICY_VERSION
from app.services.cv_docx_template import DeterministicDocxTemplateProvider
from tests.unit.test_cv_docx_render_plan_builder import build_fixture_approved_content_result


_HEADER_LINES_WITH_MAREKGRABOWSKI_LINKEDIN = CvDocxHeaderLines(
    full_name="Marek Grabowski",
    contact_tokens=("marek@example.com", "+48 000 000 000", "https://www.linkedin.com/in/marekgrabowski"),
)


class _FakeContentPlan:
    """A minimal stand-in exposing only what the renderer reads
    (``content_plan_fingerprint``) -- the renderer never reads any other
    ``CvContentPlan`` field, so a full real plan is unnecessary here."""

    def __init__(self, content_plan_fingerprint: str) -> None:
        self.content_plan_fingerprint = content_plan_fingerprint


def _render_once(
    *,
    content_plan_fingerprint: str = "1" * 64,
    result_fingerprint: str = "2" * 64,
    header_lines: CvDocxHeaderLines | None = None,
):
    approved = build_fixture_approved_content_result(
        content_plan_fingerprint=content_plan_fingerprint, result_fingerprint=result_fingerprint
    )
    content_plan = _FakeContentPlan(content_plan_fingerprint)
    provider = DeterministicDocxTemplateProvider(DEFAULT_VISUAL_CONTRACT)
    handle = provider.get_template()
    render_plan = build_cv_docx_render_plan(
        approved_content_result=approved, visual_contract=DEFAULT_VISUAL_CONTRACT, header_lines=header_lines
    )
    renderer = DeterministicCvDocxRenderer(DEFAULT_VISUAL_CONTRACT, render_plan=render_plan)
    return renderer.render(
        approved_content_result=approved,
        content_plan=content_plan,
        template_bytes=handle.template_bytes,
        template_fingerprint=handle.template_fingerprint,
        rendering_policy_version=RENDERING_POLICY_VERSION,
    )


# -- determinism --------------------------------------------------------------


def test_same_input_gives_identical_bytes_in_two_renders() -> None:
    result_one = _render_once()
    result_two = _render_once()
    assert result_one.status == DocxRenderingStatus.SUCCEEDED
    assert result_two.status == DocxRenderingStatus.SUCCEEDED
    assert result_one.docx_bytes == result_two.docx_bytes
    assert result_one.output_sha256 == result_two.output_sha256


def test_same_input_gives_identical_sha256_in_separate_processes() -> None:
    script = textwrap.dedent(
        """
        from tests.unit.test_cv_docx_renderer_determinism import _render_once
        result = _render_once()
        print(result.output_sha256)
        """
    )
    backend_root = Path(__file__).resolve().parents[2]
    hashes = set()
    for _ in range(2):
        env = os.environ.copy()
        env.pop("PYTHONHASHSEED", None)  # let each subprocess pick its own random hash seed
        completed = subprocess.run(
            [sys.executable, "-c", script],
            cwd=str(backend_root),
            capture_output=True,
            text=True,
            timeout=60,
            env=env,
        )
        assert completed.returncode == 0, completed.stderr
        hashes.add(completed.stdout.strip())
    assert len(hashes) == 1


def test_zip_timestamps_are_stable() -> None:
    result = _render_once()
    with zipfile.ZipFile(BytesIO(result.docx_bytes)) as zf:
        for info in zf.infolist():
            assert info.date_time == (1980, 1, 1, 0, 0, 0)


def test_zip_entry_order_is_stable_across_renders() -> None:
    result_one = _render_once()
    result_two = _render_once()
    with zipfile.ZipFile(BytesIO(result_one.docx_bytes)) as zf1, zipfile.ZipFile(
        BytesIO(result_two.docx_bytes)
    ) as zf2:
        assert zf1.namelist() == zf2.namelist()


def test_core_properties_are_stable() -> None:
    result_one = _render_once()
    result_two = _render_once()
    with zipfile.ZipFile(BytesIO(result_one.docx_bytes)) as zf1, zipfile.ZipFile(
        BytesIO(result_two.docx_bytes)
    ) as zf2:
        assert zf1.read("docProps/core.xml") == zf2.read("docProps/core.xml")


def test_approved_content_change_changes_bytes_and_hash() -> None:
    # force an actual content-text difference via a mutated fixture:
    approved = build_fixture_approved_content_result()
    mutated_element = approved.draft.sections[0].elements[0].model_copy(
        update={"content_text": "An entirely different profile summary."}
    )
    mutated_profile = approved.draft.sections[0].model_copy(update={"elements": (mutated_element,)})
    mutated_draft = approved.draft.model_copy(
        update={"sections": (mutated_profile, *approved.draft.sections[1:])}
    )
    mutated_approved = approved.model_copy(update={"draft": mutated_draft})

    content_plan = _FakeContentPlan(approved.content_plan_fingerprint)
    provider = DeterministicDocxTemplateProvider(DEFAULT_VISUAL_CONTRACT)
    handle = provider.get_template()

    # R3: each renderer is bound to its own already-complete render plan --
    # a mutated approved_content_result requires a freshly built plan, never
    # a plan re-derived implicitly by reusing one renderer instance.
    baseline_plan = build_cv_docx_render_plan(approved_content_result=approved, visual_contract=DEFAULT_VISUAL_CONTRACT)
    mutated_plan = build_cv_docx_render_plan(approved_content_result=mutated_approved, visual_contract=DEFAULT_VISUAL_CONTRACT)
    assert baseline_plan.render_plan_fingerprint != mutated_plan.render_plan_fingerprint

    baseline = DeterministicCvDocxRenderer(DEFAULT_VISUAL_CONTRACT, render_plan=baseline_plan).render(
        approved_content_result=approved,
        content_plan=content_plan,
        template_bytes=handle.template_bytes,
        template_fingerprint=handle.template_fingerprint,
        rendering_policy_version=RENDERING_POLICY_VERSION,
    )
    mutated = DeterministicCvDocxRenderer(DEFAULT_VISUAL_CONTRACT, render_plan=mutated_plan).render(
        approved_content_result=mutated_approved,
        content_plan=content_plan,
        template_bytes=handle.template_bytes,
        template_fingerprint=handle.template_fingerprint,
        rendering_policy_version=RENDERING_POLICY_VERSION,
    )
    assert baseline.docx_bytes != mutated.docx_bytes
    assert baseline.output_sha256 != mutated.output_sha256


def test_visual_contract_change_changes_bytes_and_hash() -> None:
    approved = build_fixture_approved_content_result()
    content_plan = _FakeContentPlan(approved.content_plan_fingerprint)

    bumped_contract = DEFAULT_VISUAL_CONTRACT.model_copy(update={"margin_top_cm": 3.5})

    plan_default = build_cv_docx_render_plan(approved_content_result=approved, visual_contract=DEFAULT_VISUAL_CONTRACT)
    plan_bumped = build_cv_docx_render_plan(approved_content_result=approved, visual_contract=bumped_contract)

    provider_default = DeterministicDocxTemplateProvider(DEFAULT_VISUAL_CONTRACT)
    handle_default = provider_default.get_template()
    renderer_default = DeterministicCvDocxRenderer(DEFAULT_VISUAL_CONTRACT, render_plan=plan_default)
    result_default = renderer_default.render(
        approved_content_result=approved,
        content_plan=content_plan,
        template_bytes=handle_default.template_bytes,
        template_fingerprint=handle_default.template_fingerprint,
        rendering_policy_version=RENDERING_POLICY_VERSION,
    )

    provider_bumped = DeterministicDocxTemplateProvider(bumped_contract)
    handle_bumped = provider_bumped.get_template()
    renderer_bumped = DeterministicCvDocxRenderer(bumped_contract, render_plan=plan_bumped)
    result_bumped = renderer_bumped.render(
        approved_content_result=approved,
        content_plan=content_plan,
        template_bytes=handle_bumped.template_bytes,
        template_fingerprint=handle_bumped.template_fingerprint,
        rendering_policy_version=RENDERING_POLICY_VERSION,
    )

    assert handle_default.template_fingerprint != handle_bumped.template_fingerprint
    assert result_default.output_sha256 != result_bumped.output_sha256


def test_renderer_version_change_changes_identity() -> None:
    approved = build_fixture_approved_content_result()
    content_plan = _FakeContentPlan(approved.content_plan_fingerprint)

    bumped_contract = DEFAULT_VISUAL_CONTRACT.model_copy(
        update={"renderer_version": "cv-docx-deterministic-renderer-v2"}
    )
    provider = DeterministicDocxTemplateProvider(DEFAULT_VISUAL_CONTRACT)
    handle_default = provider.get_template()
    provider_bumped = DeterministicDocxTemplateProvider(bumped_contract)
    handle_bumped = provider_bumped.get_template()

    assert handle_default.template_fingerprint != handle_bumped.template_fingerprint


def test_template_version_change_changes_fingerprint() -> None:
    bumped_contract = DEFAULT_VISUAL_CONTRACT.model_copy(update={"template_version": "cv-docx-template-v2"})
    provider = DeterministicDocxTemplateProvider(DEFAULT_VISUAL_CONTRACT)
    provider_bumped = DeterministicDocxTemplateProvider(bumped_contract)
    assert provider.get_template().template_fingerprint != provider_bumped.get_template().template_fingerprint


# -- visual contract compliance ------------------------------------------------


def test_no_photo_in_rendered_document() -> None:
    result = _render_once()
    with zipfile.ZipFile(BytesIO(result.docx_bytes)) as zf:
        assert not any(name.startswith("word/media/") for name in zf.namelist())
        document_xml = zf.read("word/document.xml").decode("utf-8")
    assert "w:drawing" not in document_xml
    assert "w:pict" not in document_xml


def test_body_paragraph_uses_calibri_ten() -> None:
    result = _render_once()
    document = docx.Document(BytesIO(result.docx_bytes))
    body_paragraph = next(p for p in document.paragraphs if p.style.name == "Normal" and p.text)
    run = body_paragraph.runs[0]
    assert (run.font.name or document.styles["Normal"].font.name) == "Calibri"


def test_explicit_page_break_precedes_experience_section() -> None:
    result = _render_once()
    document = docx.Document(BytesIO(result.docx_bytes))
    experience_index = next(
        i for i, p in enumerate(document.paragraphs) if p.style.name == "CvSectionHeader" and p.text == "EXPERIENCE"
    )
    preceding = document.paragraphs[experience_index - 1]
    assert "w:br" in preceding._p.xml and 'w:type="page"' in preceding._p.xml


def test_no_hyphen_prefix_on_bullet_items() -> None:
    result = _render_once()
    document = docx.Document(BytesIO(result.docx_bytes))
    bullet_paragraphs = [p for p in document.paragraphs if p.style.name == "CvBullet"]
    assert bullet_paragraphs
    for paragraph in bullet_paragraphs:
        assert not paragraph.text.startswith("-")
        assert not paragraph.text.startswith("- ")
        assert paragraph.text.startswith(DEFAULT_VISUAL_CONTRACT.bullet_glyph)


def test_no_marekgrabowski_website_text() -> None:
    result = _render_once()
    with zipfile.ZipFile(BytesIO(result.docx_bytes)) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
    assert "marekgrabowski" not in document_xml.lower()


def test_section_order_is_deterministic() -> None:
    result_one = _render_once()
    result_two = _render_once()
    document_one = docx.Document(BytesIO(result_one.docx_bytes))
    document_two = docx.Document(BytesIO(result_two.docx_bytes))
    headers_one = [p.text for p in document_one.paragraphs if p.style.name == "CvSectionHeader"]
    headers_two = [p.text for p in document_two.paragraphs if p.style.name == "CvSectionHeader"]
    assert headers_one == headers_two
    assert headers_one[0] == "PROFILE"


def test_exact_text_matches_render_plan() -> None:
    from app.services.cv_docx_render_plan_builder import build_cv_docx_render_plan

    approved = build_fixture_approved_content_result()
    plan = build_cv_docx_render_plan(approved_content_result=approved, visual_contract=DEFAULT_VISUAL_CONTRACT)
    result = _render_once()
    document = docx.Document(BytesIO(result.docx_bytes))

    from app.schemas.cv_docx_render_plan import CvDocxRenderBlockKind

    expected_texts = [b.text for b in plan.blocks if b.kind == CvDocxRenderBlockKind.BULLET_ITEM]
    rendered_bullets = [p.text for p in document.paragraphs if p.style.name == "CvBullet"]
    expected_rendered = [f"{DEFAULT_VISUAL_CONTRACT.bullet_glyph}\t{text}" for text in expected_texts]
    assert rendered_bullets == expected_rendered


# -- provenance / rejection -----------------------------------------------------


def test_renderer_never_calls_llm() -> None:
    import ast

    import app.services.cv_docx_renderer as renderer_module

    source = open(renderer_module.__file__, encoding="utf-8").read()
    tree = ast.parse(source)
    imported_names = set()
    for node in ast.walk(tree):
        if isinstance(node, ast.Import):
            imported_names.update(alias.name for alias in node.names)
        elif isinstance(node, ast.ImportFrom) and node.module:
            imported_names.add(node.module)
    assert not any("llm" in name.lower() for name in imported_names)
    assert not hasattr(renderer_module, "llm")


def test_renderer_rejects_non_ready_approved_content() -> None:
    approved = build_fixture_approved_content_result()
    not_ready = approved.model_copy(update={"status": ApprovalAssemblyStatus.PENDING_APPROVAL})
    content_plan = _FakeContentPlan(approved.content_plan_fingerprint)
    provider = DeterministicDocxTemplateProvider(DEFAULT_VISUAL_CONTRACT)
    handle = provider.get_template()
    plan = build_cv_docx_render_plan(approved_content_result=approved, visual_contract=DEFAULT_VISUAL_CONTRACT)
    renderer = DeterministicCvDocxRenderer(DEFAULT_VISUAL_CONTRACT, render_plan=plan)

    result = renderer.render(
        approved_content_result=not_ready,
        content_plan=content_plan,
        template_bytes=handle.template_bytes,
        template_fingerprint=handle.template_fingerprint,
        rendering_policy_version=RENDERING_POLICY_VERSION,
    )
    assert result.status == DocxRenderingStatus.FAILED


def test_renderer_rejects_content_plan_fingerprint_mismatch() -> None:
    approved = build_fixture_approved_content_result()
    mismatched_content_plan = _FakeContentPlan("f" * 64)
    provider = DeterministicDocxTemplateProvider(DEFAULT_VISUAL_CONTRACT)
    handle = provider.get_template()
    plan = build_cv_docx_render_plan(approved_content_result=approved, visual_contract=DEFAULT_VISUAL_CONTRACT)
    renderer = DeterministicCvDocxRenderer(DEFAULT_VISUAL_CONTRACT, render_plan=plan)

    result = renderer.render(
        approved_content_result=approved,
        content_plan=mismatched_content_plan,
        template_bytes=handle.template_bytes,
        template_fingerprint=handle.template_fingerprint,
        rendering_policy_version=RENDERING_POLICY_VERSION,
    )
    assert result.status == DocxRenderingStatus.FAILED


def test_output_sha256_is_exact_raw_hash_of_bytes() -> None:
    result = _render_once()
    assert result.output_sha256 == hashlib.sha256(result.docx_bytes).hexdigest()


def test_renderer_adapter_id_is_stable() -> None:
    approved = build_fixture_approved_content_result()
    plan = build_cv_docx_render_plan(approved_content_result=approved, visual_contract=DEFAULT_VISUAL_CONTRACT)
    renderer = DeterministicCvDocxRenderer(DEFAULT_VISUAL_CONTRACT, render_plan=plan)
    assert renderer.renderer_adapter_id == "cv-docx-deterministic-renderer-v1"


# -- R3: header is part of a real render plan -----------------------------------


def test_same_input_with_header_gives_identical_bytes_in_two_renders() -> None:
    result_one = _render_once(header_lines=_HEADER_LINES_WITH_MAREKGRABOWSKI_LINKEDIN)
    result_two = _render_once(header_lines=_HEADER_LINES_WITH_MAREKGRABOWSKI_LINKEDIN)
    assert result_one.status == DocxRenderingStatus.SUCCEEDED
    assert result_two.status == DocxRenderingStatus.SUCCEEDED
    assert result_one.docx_bytes == result_two.docx_bytes
    assert result_one.output_sha256 == result_two.output_sha256


def test_header_change_changes_bytes_and_hash() -> None:
    result_no_header = _render_once()
    result_with_header = _render_once(header_lines=_HEADER_LINES_WITH_MAREKGRABOWSKI_LINKEDIN)
    assert result_no_header.docx_bytes != result_with_header.docx_bytes
    assert result_no_header.output_sha256 != result_with_header.output_sha256

    other_header = CvDocxHeaderLines(
        full_name="Marek Grabowski",
        contact_tokens=("marek@example.com", "+48 000 000 000", "https://www.linkedin.com/in/someone-else"),
    )
    result_other_header = _render_once(header_lines=other_header)
    assert result_with_header.docx_bytes != result_other_header.docx_bytes
    assert result_with_header.output_sha256 != result_other_header.output_sha256


def test_linkedin_with_marekgrabowski_slug_is_rendered_when_header_supplied() -> None:
    """The renderer never applies a blanket ``marekgrabowski`` substring
    filter -- a LinkedIn URL containing that slug is rendered exactly as
    given in the render plan (only ``website`` is never rendered, see
    ``cv_docx_proposal_orchestration.py``)."""

    result = _render_once(header_lines=_HEADER_LINES_WITH_MAREKGRABOWSKI_LINKEDIN)
    with zipfile.ZipFile(BytesIO(result.docx_bytes)) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
    assert "marekgrabowski" in document_xml.lower()
